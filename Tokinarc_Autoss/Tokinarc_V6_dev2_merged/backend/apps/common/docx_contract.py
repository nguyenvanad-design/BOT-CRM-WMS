"""
Tokinarc — apps/common/docx_contract.py
Sinh HỢP ĐỒNG MUA BÁN dạng Word (.docx): letterhead AUTOSS + quốc hiệu +
hai bên (A=AUTOSS bán, B=khách) + điều khoản + bảng hàng + ký tên.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.shared import Inches, Pt, RGBColor

from apps.common.company import COMPANY, logo_path, vnd_to_words

_FLAME = RGBColor(0xE2, 0x58, 0x22)


def _center(p):
    p.alignment = ALIGN.CENTER
    return p


def _customer_info(cust):
    addr = cust.address if isinstance(cust.address, dict) else {}
    address = ', '.join(x for x in [addr.get('street'), addr.get('district'),
                                    addr.get('city')] if x)
    pc = cust.contacts.filter(is_primary=True).first() or cust.contacts.first()
    return address or '......', (pc.phone if pc else '......'), (cust.tax_code or '......')


def build_contract_docx(contract) -> bytes:
    c = contract
    doc = Document()

    # ── Letterhead ───────────────────────────────────────────────
    lp = logo_path()
    if lp:
        try:
            doc.add_picture(lp, width=Inches(1.3))
        except Exception:
            pass
    p = doc.add_paragraph()
    run = p.add_run(COMPANY['name']); run.bold = True; run.font.size = Pt(12)
    run.font.color.rgb = _FLAME
    for line in (f"ĐC: {COMPANY['address']}",
                 f"MST: {COMPANY['tax_code']}    ĐT: {COMPANY['phone']}    Email: {COMPANY['email']}"):
        pr = doc.add_paragraph(line); pr.runs[0].font.size = Pt(9)

    # ── Quốc hiệu ────────────────────────────────────────────────
    q = _center(doc.add_paragraph()); r = q.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'); r.bold = True
    q2 = _center(doc.add_paragraph()); r2 = q2.add_run('Độc lập – Tự do – Hạnh phúc'); r2.bold = True
    _center(doc.add_paragraph('─────────────'))

    t = _center(doc.add_paragraph()); rt = t.add_run('HỢP ĐỒNG MUA BÁN HÀNG HÓA')
    rt.bold = True; rt.font.size = Pt(15); rt.font.color.rgb = _FLAME
    _center(doc.add_paragraph(f"Số: {c.code}/{c.created_at.year}/AUTOSS"))

    doc.add_paragraph('- Căn cứ Bộ luật Dân sự và Luật Thương mại hiện hành;')
    doc.add_paragraph('- Căn cứ nhu cầu và khả năng của hai bên.')
    doc.add_paragraph(f"Hôm nay, ngày {c.created_at.strftime('%d tháng %m năm %Y')}, "
                      f"chúng tôi gồm:")

    # ── Bên A / Bên B ────────────────────────────────────────────
    a = doc.add_paragraph(); ra = a.add_run('BÊN BÁN (BÊN A): '); ra.bold = True
    a.add_run(COMPANY['name'])
    doc.add_paragraph(f"   Địa chỉ: {COMPANY['address']}")
    doc.add_paragraph(f"   MST: {COMPANY['tax_code']}      Điện thoại: {COMPANY['phone']}")
    doc.add_paragraph(f"   Đại diện: Ông/Bà {COMPANY['director']}      Chức vụ: {COMPANY['director_title']}")

    addr, phone, mst = _customer_info(c.customer)
    b = doc.add_paragraph(); rb = b.add_run('BÊN MUA (BÊN B): '); rb.bold = True
    b.add_run(c.customer.name)
    doc.add_paragraph(f"   Địa chỉ: {addr}")
    doc.add_paragraph(f"   MST: {mst}      Điện thoại: {phone}")
    doc.add_paragraph("   Đại diện: Ông/Bà ......................      Chức vụ: ..............")

    doc.add_paragraph('Hai bên thống nhất ký hợp đồng với các điều khoản sau:')

    # ── Điều 1: hàng hóa ─────────────────────────────────────────
    e1 = doc.add_paragraph(); re1 = e1.add_run('ĐIỀU 1: NỘI DUNG HỢP ĐỒNG'); re1.bold = True
    doc.add_paragraph('Bên A đồng ý bán cho Bên B các mặt hàng theo bảng sau:')

    # dòng hàng: ưu tiên từ báo giá gắn HĐ; nếu không có thì 1 dòng tổng quát
    lines = []
    if getattr(c, 'quote_id', None):
        lines = [(l.part_name or l.part_no, l.qty, int(l.unit_price_vnd or 0),
                  int(l.qty * (l.unit_price_vnd or 0))) for l in c.quote.lines.all()]
    if not lines:
        lines = [('Hàng hóa/dịch vụ theo thỏa thuận', 1, int(c.value_vnd or 0), int(c.value_vnd or 0))]

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['STT', 'Tên hàng', 'SL', 'Đơn giá', 'Thành tiền']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for idx, (name, qty, price, amount) in enumerate(lines, start=1):
        cells = tbl.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = str(name)
        cells[2].text = str(qty)
        cells[3].text = f"{price:,}".replace(',', '.')
        cells[4].text = f"{amount:,}".replace(',', '.')

    total = int(c.value_vnd or 0)
    pt = doc.add_paragraph()
    pt.add_run(f"Tổng giá trị hợp đồng: ").bold = True
    pt.add_run(f"{total:,}".replace(',', '.') + ' đ')
    pw = doc.add_paragraph(); pw.add_run(f"Bằng chữ: {vnd_to_words(total)}.").italic = True

    # ── Điều 2: thanh toán ───────────────────────────────────────
    e2 = doc.add_paragraph(); e2.add_run('ĐIỀU 2: GIÁ TRỊ & PHƯƠNG THỨC THANH TOÁN').bold = True
    paid = int(c.paid_vnd or 0)
    doc.add_paragraph(f"   Tổng giá trị: {total:,}".replace(',', '.') + ' đ'
                      + f"   |   Đã thanh toán: {paid:,}".replace(',', '.') + ' đ'
                      + f"   |   Còn lại: {total - paid:,}".replace(',', '.') + ' đ')
    doc.add_paragraph('   Hình thức: chuyển khoản hoặc tiền mặt.')

    # ── Điều 3: thời hạn ─────────────────────────────────────────
    e3 = doc.add_paragraph(); e3.add_run('ĐIỀU 3: THỜI HẠN & GIAO HÀNG').bold = True
    hieu_luc = (f"{c.start_date.strftime('%d/%m/%Y')} – {c.end_date.strftime('%d/%m/%Y')}"
                if c.start_date and c.end_date else 'theo thỏa thuận hai bên')
    doc.add_paragraph(f"   Thời hạn hiệu lực: {hieu_luc}.")

    # ── Điều 4: chung ────────────────────────────────────────────
    e4 = doc.add_paragraph(); e4.add_run('ĐIỀU 4: ĐIỀU KHOẢN CHUNG').bold = True
    doc.add_paragraph('   Hai bên cam kết thực hiện đúng các điều khoản. Mọi tranh chấp '
                      'ưu tiên thương lượng; nếu không thành sẽ đưa ra Tòa án có thẩm quyền '
                      'giải quyết. Hợp đồng lập thành 02 bản, mỗi bên giữ 01 bản có giá trị như nhau.')

    # ── Chữ ký ───────────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2)
    sa = sig.rows[0].cells
    _center(sa[0].paragraphs[0]).add_run('ĐẠI DIỆN BÊN A').bold = True
    _center(sa[1].paragraphs[0]).add_run('ĐẠI DIỆN BÊN B').bold = True
    sb = sig.rows[1].cells
    _center(sb[0].paragraphs[0]).add_run('(Ký, đóng dấu)').italic = True
    _center(sb[1].paragraphs[0]).add_run('(Ký, ghi rõ họ tên)').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
