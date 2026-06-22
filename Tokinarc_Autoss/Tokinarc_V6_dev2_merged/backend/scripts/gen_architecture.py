"""
Sinh KIẾN TRÚC DỰ ÁN (Word .docx) — sơ đồ (ASCII monospace) + giải thích + bảng.
Chạy: python scripts/gen_architecture.py [output.docx]
"""
from __future__ import annotations

import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FLAME = RGBColor(0xE2, 0x5A, 0x1C)
BLUE = RGBColor(0x12, 0x4D, 0xB5)
GREY = RGBColor(0x33, 0x33, 0x33)


def main(out: str):
    d = Document()
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(10.5)

    def h(text, level=1):
        p = d.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = FLAME if level <= 1 else GREY
        return p

    def para(text):
        d.add_paragraph(text)

    def kv(label, text):
        p = d.add_paragraph()
        r = p.add_run(label + ': '); r.bold = True; r.font.color.rgb = FLAME
        p.add_run(text)

    def bullet(text):
        d.add_paragraph(text, style='List Bullet')

    def diagram(text):
        """Khối sơ đồ ASCII — font đều để giữ thẳng hàng khung."""
        for line in text.split('\n'):
            p = d.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(line if line else ' ')
            r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = BLUE
        d.add_paragraph()

    def table(headers, rows):
        t = d.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
        for i, hd in enumerate(headers):
            t.rows[0].cells[i].text = hd
            for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v)
        d.add_paragraph()

    # ── Bìa ──
    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('KIẾN TRÚC HỆ THỐNG\nTOKINARC CRM · WMS · CEO')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = FLAME
    s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run('Sơ đồ thành phần · Module backend · Luồng dữ liệu · Phân quyền · Tech stack').italic = True
    d.add_paragraph()
    para('Tài liệu mô tả kiến trúc tổng thể: các thành phần triển khai, cấu trúc '
         'module backend, dòng chảy dữ liệu nghiệp vụ (ERP), cơ chế xác thực/phân quyền '
         'và công nghệ sử dụng.')
    d.add_page_break()

    # ── 1. Thành phần ──
    h('1. Kiến trúc thành phần (deployment)', 1)
    diagram(
        "NGƯỜI DÙNG (theo role: customer/sales/warehouse/service/manager/ceo/admin)\n"
        "        │ HTTPS\n"
        "        ├──────────────────────────────┐\n"
        "        ▼                               ▼\n"
        "┌──────────────────────────┐   ┌──────────────────────────┐\n"
        "│ FRONTEND (React/Vite)    │   │ BOT KHÁCH (FastAPI :8080)│\n"
        "│ :5173/5174  SPA          │   │ catalog/FAISS + Gemini   │\n"
        "│ CRM · WMS · CEO          │   │ CÔ LẬP (ko chạm CRM/WMS) │\n"
        "│ Trợ lý nội bộ · Chuông   │   └──────────────────────────┘\n"
        "│ Quét mã (zxing)          │        ▲ /chatbot (X-API-Key)\n"
        "└──────────┬───────────────┘────────┘\n"
        "           │ /api  (JWT Bearer)\n"
        "           ▼\n"
        "┌──────────────────────────────────────────────────────────┐\n"
        "│        BACKEND — Django 5 + DRF (:8000)                   │\n"
        "│  SimpleJWT · drf-spectacular · permission theo role       │\n"
        "│  apps: accounts crm sales wms catalog analytics storage   │\n"
        "│        common (+ learning)                               │\n"
        "└───┬───────────────┬──────────────────────┬───────────────┘\n"
        "    │ ORM           │ event publish         │ files\n"
        "    ▼               ▼                       ▼\n"
        "┌──────────┐  ┌───────────────────┐  ┌──────────────┐\n"
        "│PostgreSQL│  │Postgres LISTEN/   │  │ MinIO (S3)   │\n"
        "│(dev:SQLite)│ │NOTIFY (eventbus)  │  │ ghi âm/recap │\n"
        "└──────────┘  └───────────────────┘  └──────────────┘"
    )
    kv('Frontend', 'SPA React; gọi /api (JWT) tới Django và /chatbot (proxy chèn X-API-Key) tới bot khách.')
    kv('Backend', 'Django/DRF là lõi nghiệp vụ + phân quyền + nguồn dữ liệu thật.')
    kv('Bot khách', 'Service riêng (FastAPI) chỉ tra cứu sản phẩm, KHÔNG truy cập dữ liệu nội bộ.')

    # ── 2. Module backend ──
    h('2. Module backend (Django apps)', 1)
    table(['App', 'Vai trò', 'Model/khối chính'],
          [['accounts', 'Người dùng + phân quyền', 'User, Role; roles.py (single source: hierarchy, capability, intent)'],
           ['catalog', 'Danh mục sản phẩm', 'Part (838), Torch (121), Compatibility, TorchPartMapping'],
           ['crm', 'Khách hàng & bán hàng', 'Customer, Contact, Lead, Opportunity, Quote(+Line, duyệt 2 cấp), Contract, Visit, Activity, Ticket; import; /360/; /timeline/'],
           ['sales', 'Đơn bán & thu tiền', 'SalesOrder(+Line), Payment (ký→giao→sinh phiếu xuất)'],
           ['wms', 'Kho vận', 'Warehouse>Zone>Bin(tầng); InventoryItem(FIFO), Lot(FEFO), SerialNumber; Inbound/Outbound; PickListItem; StockMovement; CycleCount; scan'],
           ['analytics', 'Báo cáo & trợ lý nội bộ', 'services (KPI/doanh thu/công nợ/tồn/forecast); assistant.py (bot nội bộ role-gated)'],
           ['storage', 'Lưu file', 'FileObject + UploadView → MinIO'],
           ['common', 'Nền tảng dùng chung', 'BaseModel(UUID7), SoftDelete, AuditLog, Notification']])

    # ── 3. Luồng dữ liệu ──
    h('3. Luồng dữ liệu xương sống (ERP)', 1)
    diagram(
        "LEAD ─convert─► KHÁCH HÀNG ─► CƠ HỘI ─► BÁO GIÁ\n"
        "                                          │ Duyệt CẤP 1 (Manager)\n"
        "                            tổng < ngưỡng │           tổng ≥ ngưỡng\n"
        "                                          ▼                    ▼\n"
        "                                     ĐÃ DUYỆT          CHỜ CEO ─Duyệt CẤP 2(CEO)─► ĐÃ DUYỆT\n"
        "                                          │ (nút Tạo đơn)\n"
        "                                          ▼\n"
        "   ĐƠN BÁN ─Ký─► ─Giao─► [TỰ SINH] PHIẾU XUẤT WMS ─quét/pick (FIFO/FEFO)─► TRỪ TỒN\n"
        "        │ Thu tiền\n"
        "        ▼\n"
        "   CÔNG NỢ ──────────────┐\n"
        "                         ▼\n"
        "WMS: ASN ─► PHIẾU NHẬP ─quét/confirm─► CỘNG TỒN (Lot/Serial)\n"
        "                         │\n"
        "        CRM + WMS + Sales ─────────────► ANALYTICS / CEO (KPI · báo cáo · AI summary)"
    )
    kv('Điểm nối CRM↔WMS', 'Bấm "Giao" đơn bán → tự tạo phiếu xuất kho (gắn mã đơn).')
    kv('Điểm hội tụ', 'Mọi số liệu CRM/WMS/Sales đổ về Analytics để CEO xem (chỉ đọc).')
    kv('Ghi vết', 'Mọi biến động tồn tạo 1 dòng StockMovement (sổ cái).')

    # ── 4. Phân quyền ──
    h('4. Xác thực & phân quyền', 1)
    diagram(
        "Đăng nhập → /auth/login/ → JWT (access + refresh)\n"
        "Mỗi request /api → Bearer token → DRF kiểm tra role:\n"
        "  • CRM/Sales : CustomerPermission (WRITE_ROLES) + ownership (sale: của mình)\n"
        "  • WMS       : WMSPermission (warehouse/manager/ceo/admin)\n"
        "  • CEO       : IsManagerOrAdmin (manager/ceo/admin)\n"
        "  • Bot nội bộ: IsInternalStaff + can_use_intent(role, intent)\n"
        "Bot khách   : chỉ X-API-Key → KHÔNG danh tính user → ko chạm dữ liệu nội bộ"
    )
    table(['Vai trò', 'CRM', 'WMS', 'Duyệt BG', 'CEO/Báo cáo'],
          [['sales', 'của mình', '–', '–', '–'],
           ['warehouse', '–', '✓', '–', '–'],
           ['service', 'ticket', '–', '–', '–'],
           ['manager', 'tất cả', '✓', 'cấp 1', '✓'],
           ['ceo', 'tất cả', '✓', 'cấp 2', '✓'],
           ['admin', 'tất cả', '✓', 'cả 2', '✓']])
    kv('Nguồn quyền', 'apps/accounts/roles.py là SINGLE SOURCE (hierarchy + capability + intent của bot).')

    # ── 5. Hai bot AI ──
    h('5. Hai "bộ não" AI (tách biệt)', 1)
    table(['', 'Bot KHÁCH (FastAPI)', 'Bot NỘI BỘ (Django assistant)'],
          [['Người dùng', 'Khách hàng', 'Nhân viên (JWT + role)'],
           ['Dữ liệu', 'Catalog / FAISS (chỉ đọc)', 'CRM/WMS/Analytics (đọc + ghi thật)'],
           ['Khả năng', 'Tra cứu sản phẩm', 'Báo giá, hợp đồng, phiếu kho, báo cáo, tra đơn/tồn'],
           ['Bảo mật', 'Cô lập, X-API-Key', 'Gate theo role; hành động tạo BẢN NHÁP']])

    # ── 6. Tech stack ──
    h('6. Công nghệ sử dụng', 1)
    table(['Tầng', 'Công nghệ'],
          [['Frontend', 'React 18 · Vite 5 · TypeScript · React Query · Zustand · React Router · Tailwind · recharts · @zxing · react-hook-form'],
           ['Backend', 'Django 5 · DRF · SimpleJWT · drf-spectacular · openpyxl/python-docx'],
           ['Bot khách', 'FastAPI · Google Gemini · FAISS · bge-m3'],
           ['Bot nội bộ', 'Django-native (analytics/assistant) · Gemini intent + fallback từ khóa'],
           ['Dữ liệu', 'PostgreSQL (prod) / SQLite (dev) · MinIO (S3) · Postgres LISTEN/NOTIFY'],
           ['Hạ tầng', 'Docker / docker-compose · nginx (reverse proxy)']])

    # ── 7. Cấu trúc kho (WMS) ──
    h('7. Cấu trúc kho (WMS) — đa kho', 1)
    diagram(
        "WAREHOUSE (HCM, HN, …)\n"
        "   └─ ZONE = nhóm sản phẩm (SUNG, MIG, TIG, THAN, CAP, ROBOT, LK, WX)\n"
        "        └─ RACK (tầng) = phân loại con (vd SUNG: T1 cầm tay / T2 TIG / T3 robot / T4 nước)\n"
        "             └─ BIN (ô) = vị trí cụ thể  (mã: HCM-SUNG-T3-B01)\n"
        "                  └─ InventoryItem (tồn) · Lot (lô/FEFO) · SerialNumber"
    )
    kv('Dựng zone', 'Lệnh build_zones tạo zone/tầng/ô theo dữ liệu sản phẩm (v20) cho từng kho.')

    d.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    main(sys.argv[1] if len(sys.argv) > 1 else 'KIEN_TRUC_DU_AN.docx')
