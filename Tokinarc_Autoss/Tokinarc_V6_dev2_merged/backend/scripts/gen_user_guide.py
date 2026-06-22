"""
Sinh HƯỚNG DẪN SỬ DỤNG CHI TIẾT (Word .docx) — mô tả từng tính năng & từng nút.
Chạy: python scripts/gen_user_guide.py [output.docx]
"""
from __future__ import annotations

import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FLAME = RGBColor(0xE2, 0x5A, 0x1C)
GREY = RGBColor(0x33, 0x33, 0x33)


def main(out: str):
    d = Document()
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(10.5)

    def h(text, level=1):
        p = d.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = FLAME if level <= 1 else GREY
        return p

    def para(text, bold=False, italic=False):
        p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
        return p

    def bullet(text):
        d.add_paragraph(text, style='List Bullet')

    def steps(items):
        for it in items:
            d.add_paragraph(it, style='List Number')

    def buttons(rows):
        """Bảng 2 cột: Nút/Thao tác | Chức năng."""
        t = d.add_table(rows=1, cols=2); t.style = 'Light Grid Accent 1'
        hc = t.rows[0].cells
        hc[0].text = 'Nút / Thao tác'; hc[1].text = 'Chức năng'
        for c in hc:
            for r in c.paragraphs[0].runs: r.bold = True
        for a, b in rows:
            cells = t.add_row().cells
            cells[0].text = a; cells[1].text = b
            for r in cells[0].paragraphs[0].runs: r.bold = True
        d.add_paragraph()

    def table(headers, rows):
        t = d.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
        for i, hd in enumerate(headers):
            t.rows[0].cells[i].text = hd
            for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v)
        d.add_paragraph()

    # ════════════════ BÌA ════════════════
    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('HƯỚNG DẪN SỬ DỤNG CHI TIẾT\nHỆ THỐNG TOKINARC CRM · WMS')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = FLAME
    s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run('Mô tả từng màn hình, từng tính năng và từng nút bấm').italic = True
    d.add_paragraph()
    para('Quy ước:', bold=True)
    bullet('“Nút” = nút bấm trên màn hình; “menu” = mục bên thanh trái.')
    bullet('“+” nghĩa là vai trò đó và cao hơn (vd Sale+ = sale, manager, ceo, admin).')
    d.add_page_break()

    # ════════════════ 1. TỔNG QUAN GIAO DIỆN ════════════════
    h('1. Tổng quan giao diện', 1)
    para('Sau khi đăng nhập, màn hình chia 3 vùng:')
    bullet('Thanh trái: chuyển phân hệ (CRM / WMS / CEO) + danh sách menu của phân hệ.')
    bullet('Thanh trên (header): chuông thông báo, tên người dùng + vai trò, nút Đăng xuất.')
    bullet('Đáy trang: thanh "Trợ lý nội bộ" (chat) — luôn hiện ở mọi màn.')
    h('1.1. Thanh trên (header) — các nút', 2)
    buttons([
        ('☰ (chỉ mobile)', 'Mở/đóng menu thanh trái trên điện thoại.'),
        ('🔔 Chuông', 'Hiện số thông báo chưa đọc. Bấm để mở danh sách.'),
        ('Đọc hết', '(trong dropdown chuông) Đánh dấu mọi thông báo là đã đọc.'),
        ('1 dòng thông báo', 'Bấm để tới màn liên quan + tự đánh dấu đã đọc.'),
        ('Đăng xuất', 'Thoát tài khoản, quay về màn đăng nhập.'),
    ])
    h('1.2. Bộ chuyển phân hệ (thanh trái trên cùng)', 2)
    buttons([
        ('CRM', 'Mở phân hệ Khách hàng/Bán hàng (mọi nhân viên).'),
        ('WMS', 'Mở phân hệ Kho (chỉ Kho/Manager/CEO/Admin thấy).'),
        ('CEO', 'Mở bảng điều hành (chỉ Manager/CEO/Admin thấy).'),
    ])
    h('1.3. Thanh Trợ lý nội bộ (chat đáy trang)', 2)
    buttons([
        ('Ô nhập + nút Hỏi', 'Gõ câu lệnh tiếng Việt rồi bấm Hỏi (hoặc Enter).'),
        ('Chip gợi ý', 'Bấm 1 gợi ý có sẵn (đổi theo vai trò) để hỏi nhanh.'),
        ('Thu gọn / Mở', 'Ẩn/hiện khung hội thoại.'),
        ('🗑 (thùng rác)', 'Xóa hội thoại hiện tại.'),
    ])

    # ════════════════ 2. ĐĂNG NHẬP ════════════════
    h('2. Đăng nhập', 1)
    steps(['Nhập Tên đăng nhập + Mật khẩu.', 'Bấm nút "Đăng nhập".',
           'Sai mật khẩu sẽ hiện thông báo lỗi; nhập lại.'])
    table(['Tài khoản', 'Mật khẩu', 'Vai trò'],
          [['admin', 'admin12345', 'Toàn quyền'], ['ceo1', 'tokinarc123', 'CEO'],
           ['quanly1', 'tokinarc123', 'Quản lý'], ['sale1', 'tokinarc123', 'Sale'],
           ['kho1', 'tokinarc123', 'Kho'], ['kysu1', 'tokinarc123', 'Dịch vụ']])

    # ════════════════ 3. CRM ════════════════
    h('3. Phân hệ CRM — chi tiết từng màn', 1)

    h('3.1. Dashboard (Tổng quan)', 2)
    para('Hiển thị thẻ KPI (doanh số, cơ hội, công nợ…). Chủ yếu để xem, không có nút thao tác.')

    h('3.2. Khách hàng', 2)
    buttons([
        ('Ô tìm kiếm', 'Gõ tên/mã KH để lọc (tự lọc sau ~0,3s).'),
        ('Import (Quản lý+)', 'Mở hộp thoại nhập KH cũ từ Excel/CSV (xem 3.13).'),
        ('Thêm KH', 'Mở form tạo khách hàng mới.'),
        ('Bấm vào 1 dòng KH', 'Mở hồ sơ Khách hàng 360.'),
        ('Trước / Sau', 'Chuyển trang danh sách.'),
    ])
    para('Form "Thêm/Sửa KH" — các nút:', bold=True)
    buttons([
        ('Mã KH *', 'Bắt buộc, phải bắt đầu bằng "KH" (vd KH-0001).'),
        ('Tên, Phân khúc, Vùng, Mã số thuế, Ghi chú', 'Thông tin KH.'),
        ('Thêm người liên hệ', 'Thêm dòng liên hệ (tên, SĐT, email, kênh, "chính").'),
        ('Hủy / Lưu', 'Đóng không lưu / Lưu khách hàng.'),
    ])

    h('3.3. Khách hàng 360 (hồ sơ KH)', 2)
    buttons([
        ('← Quay lại', 'Về danh sách khách hàng.'),
        ('Sửa', 'Mở form sửa thông tin KH.'),
        ('Thẻ KPI', 'Đơn đang mở, Công nợ, Ticket mở, Hoạt động gần nhất.'),
        ('Mục Cơ hội / Báo giá', 'Danh sách cơ hội & báo giá của KH (link "Tất cả…").'),
        ('Lịch sử làm việc', 'Dòng thời gian: viếng thăm, gọi, báo giá, đơn, ticket.'),
        ('🎧 Nghe ghi âm / 📄 File recap', 'Tải/nghe file đính kèm của buổi gặp/gọi.'),
    ])

    h('3.4. Leads', 2)
    buttons([
        ('Ô tìm kiếm', 'Lọc theo tên/công ty.'),
        ('Import (Quản lý+)', 'Nhập Lead cũ từ Excel/CSV.'),
        ('Tạo Lead', 'Mở form tạo lead mới.'),
        ('Bấm 1 dòng', 'Mở form sửa lead.'),
        ('Chuyển KH', 'Chuyển lead đủ điều kiện thành Khách hàng thật.'),
        ('Trước / Sau', 'Phân trang.'),
    ])

    h('3.5. Opportunity (Cơ hội)', 2)
    buttons([
        ('Tạo Opportunity', 'Mở form tạo cơ hội (KH, giá trị, xác suất, giai đoạn).'),
        ('Bấm 1 dòng', 'Mở trang chi tiết cơ hội + dòng thời gian tư vấn.'),
        ('Trước / Sau', 'Phân trang.'),
    ])
    para('Trang chi tiết Cơ hội: xem timeline (hoạt động + viếng thăm gắn cơ hội), sửa, đổi giai đoạn.')

    h('3.6. Pipeline', 2)
    para('Bảng Kanban theo giai đoạn (prospect → won/lost). Kéo/bấm để theo dõi tiến độ.')

    h('3.7. Báo giá — quan trọng (DUYỆT 2 CẤP)', 2)
    buttons([
        ('Tạo BG', 'Mở form tạo báo giá (chọn KH + dòng hàng).'),
        ('Bấm dòng (khi Nháp)', 'Sửa báo giá (chỉ sửa khi trạng thái Nháp).'),
        ('Duyệt / Duyệt (cấp 1)', '(Manager+) Duyệt báo giá. Dưới ngưỡng → Đã duyệt; ≥ ngưỡng → Chờ CEO.'),
        ('Duyệt cấp 2 (CEO)', '(CEO/Admin) Duyệt báo giá đang "Chờ CEO duyệt".'),
        ('Từ chối', '(Manager+) Từ chối, nhập lý do (hiện hộp nhập lý do).'),
        ('Tạo đơn', '(Báo giá đã duyệt) Sinh Đơn bán thật từ báo giá.'),
        ('Tạo HĐ', '(Báo giá đã duyệt) Sinh mã hợp đồng.'),
        ('Chờ duyệt / Chờ CEO duyệt', 'Nhãn trạng thái khi bạn không có quyền duyệt.'),
    ])
    para('Form "Tạo BG" — các nút:', bold=True)
    buttons([
        ('Chọn khách hàng', 'Bắt buộc.'),
        ('Mã part / Tên part / SL / Đơn giá', 'Nhập từng dòng hàng (tổng tự tính ở server).'),
        ('Thêm dòng / Xóa dòng', 'Thêm/bớt dòng hàng.'),
        ('Hủy / Tạo', 'Đóng / Lưu báo giá (trạng thái Nháp).'),
    ])

    h('3.8. Hợp đồng', 2)
    buttons([
        ('Thẻ KPI', 'Hiệu lực / Chờ ký / Hết hạn / Tổng giá trị.'),
        ('Import (Quản lý+)', 'Nhập hợp đồng cũ (cột customer_code = mã KH).'),
        ('Tạo HĐ', 'Mở form tạo hợp đồng.'),
        ('Bấm 1 dòng', 'Sửa hợp đồng.'),
    ])

    h('3.9. Công nợ (Receivables)', 2)
    buttons([
        ('Import đơn cũ (Quản lý+)', 'Nhập đơn hàng cũ từ Excel/CSV.'),
        ('Thẻ KPI', 'Tổng phải thu / Trong hạn / Quá hạn / Quá >60 ngày.'),
        ('Phân tích tuổi nợ', 'Thanh tỷ lệ theo nhóm tuổi nợ.'),
        ('Bấm 1 dòng', 'Mở hồ sơ KH liên quan.'),
    ])

    h('3.10. Visit (Viếng thăm)', 2)
    para('Form ghi viếng thăm — các nút:', bold=True)
    buttons([
        ('Khách hàng * / Ngày thăm *', 'Bắt buộc.'),
        ('Gắn cơ hội', 'Liên kết buổi thăm với 1 cơ hội (tùy chọn).'),
        ('Mục đích / Tóm tắt / Hành động tiếp theo', 'Nội dung buổi thăm.'),
        ('File ghi âm — Chọn file', 'Tải audio buổi gặp lên (MinIO).'),
        ('File recap (Word/PDF) — Chọn file', 'Tải file recap.'),
        ('Recap (văn bản)', 'Gõ tóm tắt buổi gặp.'),
        ('✕ trên file', 'Gỡ file đã chọn.'),
        ('Hủy / Lưu', 'Đóng / Lưu viếng thăm.'),
    ])

    h('3.11. Hoạt động (Activity)', 2)
    para('Tương tự Visit nhưng cho Gọi điện / Email / Zalo / Gặp mặt:')
    buttons([
        ('Loại', 'Chọn Gọi điện / Email / Zalo / Gặp mặt / Khác.'),
        ('Nội dung', 'Nội dung trao đổi.'),
        ('File ghi âm / File recap / Recap (văn bản)', 'Như mục 3.10.'),
        ('Hủy / Lưu', 'Đóng / Lưu.'),
    ])

    h('3.12. Service Ticket & Bảo hành', 2)
    buttons([
        ('Tạo Ticket', 'Mở form tạo yêu cầu hỗ trợ/bảo hành.'),
        ('Số serial', 'Gắn serial sản phẩm lỗi (liên kết WMS).'),
        ('Menu Bảo hành', 'Tra serial súng hàn + trạng thái còn/hết hạn bảo hành.'),
    ])

    h('3.13. Import dữ liệu cũ (hộp thoại Import)', 2)
    buttons([
        ('Tải file mẫu (Excel)', 'Tải file .xlsx mẫu đúng cột để điền dữ liệu cũ.'),
        ('Chọn file Excel/CSV', 'Chọn file đã điền (.xlsx hoặc .csv).'),
        ('Xem trước', 'Kiểm tra: sẽ tạo bao nhiêu / bỏ qua trùng / lỗi từng dòng (CHƯA ghi).'),
        ('Import', 'Ghi dữ liệu vào hệ thống (bỏ qua bản ghi trùng mã).'),
        ('Đóng', 'Đóng hộp thoại.'),
    ])

    h('3.14. Sản phẩm & AI Gợi ý', 2)
    bullet('Sản phẩm: 2 tab "Phụ tùng" / "Súng hàn" + ô tìm kiếm + phân trang.')
    bullet('AI Gợi ý: trang giới thiệu các trợ lý AI; bấm để mở chức năng tương ứng.')

    # ════════════════ 4. WMS ════════════════
    h('4. Phân hệ WMS — chi tiết từng màn', 1)

    h('4.1. Dashboard / Tồn kho / Serial / Lịch sử kho', 2)
    buttons([
        ('Menu Tồn kho — ô tìm kiếm', 'Lọc tồn theo tên/mã/ô.'),
        ('Menu Sắp hết hàng', 'Lọc mặt hàng tồn ≤ mức tối thiểu.'),
        ('Menu Serial', 'Danh sách serial súng hàn + trạng thái.'),
        ('Menu Lịch sử kho', 'Sổ cái biến động tồn (mọi nhập/xuất/điều chỉnh).'),
    ])

    h('4.2. Nhập kho', 2)
    buttons([
        ('Tạo đơn nhập', 'Mở form tạo đơn nhập (chọn kho + dòng hàng + ô đích).'),
        ('Quét', '(trên 1 đơn nháp/đã xác nhận) Mở cửa sổ quét nhận theo phiếu.'),
        ('Xác nhận', 'Cộng tồn theo số đã nhận của đơn.'),
    ])
    para('Cửa sổ "Quét nhận":', bold=True)
    steps(['Quét/nhập "Mã hàng" + "SL" → bấm nút quét → cộng dồn số đã nhận.',
           'Bảng hiện tiến độ từng dòng (đã/cần, "Đủ"/"Còn N").',
           'Khi đủ, bấm "Xác nhận nhận" để cộng tồn kho.'])

    h('4.3. Xuất kho', 2)
    buttons([
        ('Tạo đơn xuất', 'Mở form tạo đơn xuất.'),
        ('Pick-list', 'Xem gợi ý lấy hàng ở ô nào (FIFO/FEFO).'),
        ('Quét', 'Mở cửa sổ quét soạn hàng theo phiếu.'),
        ('Giao', 'Xác nhận giao — trừ tồn, ghi xuất kho.'),
    ])
    para('Cửa sổ "Quét soạn":', bold=True)
    steps(['Quét/nhập "Mã hàng" + "Mã ô" + "SL" → bấm nút quét → trừ tồn + cộng đã soạn.',
           'Bảng hiện tiến độ từng dòng.',
           'Khi đủ, bấm "Giao hàng".'])

    h('4.4. Quét mã (camera điện thoại)', 2)
    buttons([
        ('Tab Tra cứu', 'Quét → xem phụ tùng/serial khớp.'),
        ('Tab Nhập kho', 'Quét mã + ô + SL → cộng tồn.'),
        ('Tab Xuất kho', 'Quét mã + ô + SL → trừ tồn.'),
        ('Tab Kiểm kê', 'Quét mã + ô + số đếm → đặt lại tồn.'),
        ('Bắt đầu quét / Dừng', 'Bật/tắt camera.'),
        ('Đang quét vào: Mã hàng / Ô', 'Chọn camera điền vào ô nào.'),
        ('Nhập kho / Xuất kho / Cập nhật tồn', 'Nút thực hiện theo tab đang chọn.'),
        ('Ô nhập tay + Tra', '(tab Tra cứu) Nhập mã thủ công để tìm.'),
    ])

    h('4.5. Kiểm kê (phiên)', 2)
    buttons([
        ('Phiên mới', 'Tạo 1 phiên kiểm kê cho kho.'),
        ('Bấm 1 phiên', 'Mở phiên để quét đếm.'),
        ('Mã hàng / Mã ô / Số đếm + nút quét', 'Ghi số đếm thực tế của 1 mặt hàng tại 1 ô.'),
        ('Bảng chênh lệch', 'Hệ thống vs Đếm; cột Chênh tô xanh (+)/đỏ (−).'),
        ('Áp dụng', 'Điều chỉnh tồn theo số đếm (ghi vết). Phiên chuyển "applied".'),
        ('← Danh sách phiên', 'Quay lại danh sách.'),
    ])

    h('4.6. Kho & vị trí / Bản đồ kho / Báo cáo', 2)
    bullet('Kho & vị trí: danh sách kho, zone, ô (bin).')
    bullet('Bản đồ kho: sơ đồ trực quan các ô.')
    bullet('Báo cáo: biểu đồ tồn/biến động.')

    # ════════════════ 5. CEO ════════════════
    h('5. Phân hệ CEO (Manager/CEO/Admin)', 1)
    buttons([
        ('Bảng điều hành (Overview)', 'KPI tổng hợp toàn công ty.'),
        ('Doanh thu / Công nợ / Forecast / Tồn kho', 'Các báo cáo chuyên đề.'),
        ('AI Summary — Làm mới', 'Gọi AI tổng hợp lại số liệu mới nhất.'),
        ('AI Summary — Tải Excel', 'Xuất báo cáo điều hành ra file .xlsx.'),
    ])

    # ════════════════ 6. TRỢ LÝ NỘI BỘ ════════════════
    h('6. Trợ lý nội bộ (chat) — các câu lệnh', 1)
    para('Gõ vào thanh chat đáy trang. Trợ lý làm theo đúng quyền của bạn:')
    table(['Bạn gõ', 'Kết quả', 'Vai trò'],
          [['làm báo giá cho Công ty ABC: 5 x 001002', 'Tạo báo giá nháp', 'Sale+'],
           ['soạn hợp đồng từ báo giá BG-0007', 'Tạo hợp đồng nháp', 'Sale+'],
           ['đơn của Công ty ABC', 'Liệt kê đơn + công nợ KH', 'Sale+'],
           ['tồn 001002', 'Tra tồn theo ô', 'Mọi NV'],
           ['tra cứu phụ tùng 001002', 'Spec + giá + súng tương thích', 'Mọi NV'],
           ['nhập kho 100 x 001002', 'Lập phiếu nhập nháp', 'Kho/CEO/Admin'],
           ['xuất kho 20 x 001002', 'Lập phiếu xuất nháp', 'Kho/CEO/Admin'],
           ['báo cáo điều hành', 'Tóm tắt toàn công ty', 'Manager/CEO+'],
           ['đánh giá kế hoạch pipeline', 'Dự báo doanh thu', 'Manager/CEO+']])
    bullet('Hành động ghi tạo BẢN NHÁP; vào màn tương ứng để gửi/duyệt/xác nhận.')
    bullet('Khách hàng không dùng được trợ lý nội bộ.')

    # ════════════════ 7. THÔNG BÁO ════════════════
    h('7. Thông báo (chuông 🔔)', 1)
    bullet('Badge đỏ = số thông báo chưa đọc (tự cập nhật ~30 giây).')
    bullet('Nhận khi: báo giá chờ CEO duyệt, báo giá được duyệt/từ chối.')
    bullet('Bấm 1 thông báo → mở màn liên quan + đánh dấu đã đọc.')
    bullet('"Đọc hết" → đánh dấu tất cả đã đọc.')

    # ════════════════ 8. QUY TRÌNH MẪU ════════════════
    h('8. Quy trình mẫu (đầu–cuối)', 1)
    para('Bán hàng:', bold=True)
    steps(['Tạo Lead → Chuyển KH → tạo Cơ hội.',
           'Tạo Báo giá → Manager Duyệt (→ CEO duyệt nếu lớn).',
           'Báo giá đã duyệt → "Tạo đơn" → Ký đơn → Giao (tự sinh phiếu xuất WMS).',
           'Kho "Quét soạn" → "Giao hàng" (trừ tồn).',
           'Ghi nhận thanh toán → công nợ giảm.'])
    para('Nhập hàng về kho:', bold=True)
    steps(['Tạo đơn nhập (hoặc từ ASN).', 'Bấm "Quét" → quét từng mã + SL.',
           'Bấm "Xác nhận nhận" → cộng tồn.'])
    para('Kiểm kê định kỳ:', bold=True)
    steps(['WMS → Kiểm kê → "Phiên mới".', 'Quét đếm từng ô.',
           'Xem chênh lệch → "Áp dụng".'])

    # ════════════════ 9. FAQ ════════════════
    h('9. Câu hỏi thường gặp', 1)
    para('Không thấy menu WMS / CEO?', bold=True)
    para('WMS chỉ cho Kho/Manager/CEO/Admin; CEO chỉ cho Manager/CEO/Admin.')
    para('Không có nút Duyệt báo giá?', bold=True)
    para('Chỉ Manager/CEO/Admin được duyệt; người tạo không tự duyệt (trừ admin).')
    para('Không có nút Import?', bold=True)
    para('Chỉ Quản lý/CEO/Admin mới thấy nút Import.')
    para('Camera quét không mở được?', bold=True)
    para('Dùng ô nhập tay; trên điện thoại cần truy cập qua HTTPS và cấp quyền camera.')
    para('Quét xuất báo "Tồn khả dụng < N"?', bold=True)
    para('Ô đó không đủ hàng — chọn ô khác hoặc kiểm tra lại tồn.')

    d.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    main(sys.argv[1] if len(sys.argv) > 1 else 'HUONG_DAN_SU_DUNG_TOKINARC.docx')
