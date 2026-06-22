"""
Sinh HƯỚNG DẪN SỬ DỤNG & LUỒNG NGHIỆP VỤ (Word .docx).
Giải thích từng tính năng CRM/WMS + kịch bản luồng đầu–cuối (ai làm gì, chuyển cho ai).
Chạy: python scripts/gen_huong_dan.py [output.docx]
"""
from __future__ import annotations

import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FLAME = RGBColor(0xE2, 0x5A, 0x1C)
BLUE = RGBColor(0x12, 0x4D, 0xB5)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
GREY = RGBColor(0x33, 0x33, 0x33)


def main(out: str):
    d = Document()
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(10.5)

    def h(text, level=1):
        p = d.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = FLAME if level <= 1 else GREY
        return p

    def para(text):
        d.add_paragraph(text)

    def kv(label, text):
        p = d.add_paragraph()
        r = p.add_run(label + ': '); r.bold = True; r.font.color.rgb = FLAME
        p.add_run(text)

    def bullet(text):
        d.add_paragraph(text, style='List Bullet')

    def steps(items):
        for it in items:
            d.add_paragraph(it, style='List Number')

    def buttons(rows):
        t = d.add_table(rows=1, cols=2); t.style = 'Light Grid Accent 1'
        hc = t.rows[0].cells; hc[0].text = 'Nút / Trường'; hc[1].text = 'Tác dụng'
        for c in hc:
            for r in c.paragraphs[0].runs: r.bold = True
        for a, b in rows:
            cells = t.add_row().cells; cells[0].text = a; cells[1].text = b
            for r in cells[0].paragraphs[0].runs: r.bold = True
        d.add_paragraph()

    def flow(text):
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text); r.bold = True; r.font.name = 'Consolas'
        r.font.size = Pt(9); r.font.color.rgb = BLUE

    def table(headers, rows, widths=None):
        t = d.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
        for i, hd in enumerate(headers):
            t.rows[0].cells[i].text = hd
            for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v)
        d.add_paragraph()

    def scenario(title, steps):
        """steps = list (Bước, Người làm, Thao tác, Kết quả/Trạng thái, Chuyển cho)."""
        hp = d.add_paragraph(); r = hp.add_run('▸ ' + title)
        r.bold = True; r.font.size = Pt(12); r.font.color.rgb = GREEN
        table(['#', 'Người làm', 'Thao tác', 'Kết quả / Trạng thái', 'Chuyển cho'], steps)

    # ════════ BÌA ════════
    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('HƯỚNG DẪN SỬ DỤNG & LUỒNG NGHIỆP VỤ\nTOKINARC CRM · WMS')
    r.bold = True; r.font.size = Pt(21); r.font.color.rgb = FLAME
    s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run('Giải thích từng tính năng + kịch bản công việc từ đầu đến cuối').italic = True
    d.add_paragraph()
    para('Tài liệu gồm 2 phần: (1) Giải thích từng tính năng CRM/WMS; '
         '(2) Các LUỒNG NGHIỆP VỤ đầu–cuối — ai làm gì, hệ thống đổi trạng thái ra sao, '
         'rồi chuyển việc cho ai.')
    d.add_page_break()

    # ════════ A. TỔNG QUAN ════════
    h('A. Tổng quan & vai trò', 1)
    table(['Vai trò', 'Phụ trách chính'],
          [['Sale', 'Khách hàng, lead, cơ hội, báo giá, hợp đồng, đơn bán'],
           ['Manager', 'Duyệt báo giá cấp 1, ký đơn, thu tiền, xem báo cáo'],
           ['CEO', 'Duyệt báo giá cấp 2 (giá trị lớn), dashboard điều hành'],
           ['Nhân viên kho', 'Nhập/xuất/chuyển/quét/đếm (không sửa tồn)'],
           ['Quản lý kho', 'Điều chỉnh tồn, duyệt kiểm kê, FIFO/FEFO, KPI kho'],
           ['Dịch vụ', 'Ticket bảo hành/hỗ trợ'],
           ['Admin', 'Toàn quyền']])
    kv('Đăng nhập', 'Vào địa chỉ hệ thống → nhập tài khoản/mật khẩu → bấm "Đăng nhập".')

    # ════════ B. CRM — TỪNG TÍNH NĂNG ════════
    h('B. CRM — giải thích từng tính năng', 1)

    feats_crm = [
        ('Khách hàng', 'Lưu hồ sơ khách + 360 (KPI, liên hệ, cơ hội, báo giá, lịch sử làm việc).',
         'Thêm KH / Sửa / Import Excel / bấm vào KH để xem 360.'),
        ('Lead', 'Khách tiềm năng chưa chính thức; sàng lọc rồi "Chuyển KH".',
         'Tạo Lead → Chuyển KH (tạo Khách hàng thật).'),
        ('Cơ hội (Opportunity)', 'Theo dõi 1 thương vụ + xác suất; nuôi Pipeline/Forecast.',
         'Tạo Opportunity, đổi giai đoạn (prospect→…→won/lost).'),
        ('Báo giá (Quote)', 'Gửi giá cho khách; DUYỆT 2 CẤP theo giá trị.',
         'Tạo BG (chọn KH + dòng hàng); Duyệt cấp 1/2; Từ chối; Tạo đơn; Tạo HĐ.'),
        ('Hợp đồng', 'Văn bản cam kết; sinh từ báo giá đã duyệt hoặc tạo trực tiếp.',
         'Tạo HĐ / Import; nháp→chờ ký→hiệu lực→hết hạn.'),
        ('Đơn bán', 'Đơn giao + thu tiền; khi Giao tự sinh phiếu xuất kho.',
         'Ký → Giao → (thu tiền qua công nợ).'),
        ('Công nợ', 'Phải thu + phân tích tuổi nợ.', 'Xem; Import đơn cũ; ghi thanh toán → giảm nợ.'),
        ('Viếng thăm / Hoạt động', 'Ghi lại gặp/gọi + tải ghi âm + recap.',
         'Tạo Visit/Activity; đính file ghi âm + file recap + nhập recap chữ.'),
        ('Lịch sử làm việc', 'Dòng thời gian mọi tương tác của 1 khách.',
         'Trong hồ sơ KH 360; nghe ghi âm / tải recap.'),
        ('Ticket & Bảo hành', 'Yêu cầu hỗ trợ/bảo hành; gắn serial sản phẩm.',
         'Tạo Ticket (gắn serial); tra Bảo hành theo serial.'),
        ('Import dữ liệu cũ', 'Đưa KH/Lead/Hợp đồng/Đơn cũ vào hệ thống.',
         'Tải mẫu → điền Excel → Xem trước → Import.'),
    ]
    for name, purpose, how in feats_crm:
        h(name, 3); kv('Chức năng', purpose); kv('Thao tác', how)

    # ════════ C. WMS — TỪNG TÍNH NĂNG ════════
    h('C. WMS — giải thích từng tính năng', 1)
    feats_wms = [
        ('Tồn kho', 'Xem còn bao nhiêu, ở ô nào; lọc sắp hết hàng.', 'Menu Tồn kho / Sắp hết hàng.'),
        ('Cấu trúc kho', 'Kho > Zone (nhóm SP) > Tầng (loại con) > Ô (vị trí).', 'Kho & vị trí / Bản đồ kho.'),
        ('Nhập kho', 'Nhận hàng NCC → cộng tồn.', 'Tạo đơn nhập → Quét nhận / Xác nhận nhận.'),
        ('Xuất kho', 'Soạn & giao hàng → trừ tồn.', 'Tạo/auto từ đơn bán → Pick-list/Quét soạn → Giao.'),
        ('Quét mã', 'Quét barcode điện thoại: tra cứu / nhập / xuất / kiểm kê.', 'Menu Quét mã, chọn chế độ.'),
        ('Kiểm kê (phiên)', 'Đếm thực tế → đối chiếu → áp dụng điều chỉnh tồn.', 'Phiên mới → quét đếm → Áp dụng (Quản lý kho).'),
        ('Lô hàng (FEFO)', 'Theo dõi lô + hạn dùng; xuất lô hết hạn trước.', 'Menu Lô hàng; cảnh báo sắp hết hạn.'),
        ('Serial', 'Theo dõi từng súng hàn + bảo hành.', 'Menu Serial; lịch sử serial.'),
        ('KPI vận hành', 'Năng suất nhập/xuất, độ chính xác kiểm kê, tồn theo zone, hiệu suất NV.', 'Menu KPI vận hành (Quản lý kho+).'),
    ]
    for name, purpose, how in feats_wms:
        h(name, 3); kv('Chức năng', purpose); kv('Thao tác', how)

    # ════════ D. CÁC LUỒNG NGHIỆP VỤ (TRỌNG TÂM) ════════
    d.add_page_break()
    h('D. Các luồng nghiệp vụ đầu–cuối', 1)
    para('Phần quan trọng nhất: mô tả công việc chạy qua nhiều người như thế nào.')

    h('D.1. Luồng BÁN HÀNG đầy đủ (báo giá → giao hàng → thu tiền)', 2)
    flow("SALE tạo báo giá ─► MANAGER duyệt cấp 1 ─► (lớn) CEO duyệt cấp 2 ─►")
    flow("   SALE báo khách ─► khách ĐỒNG Ý ─► SALE tạo hợp đồng + đơn bán ─►")
    flow("   MANAGER ký đơn ─► KHO xuất hàng ─► giao khách ─► MANAGER thu tiền")
    scenario('Kịch bản chi tiết', [
        ['1', 'Sale', 'Tạo báo giá (chọn KH + dòng hàng), bấm Lưu', 'Báo giá: Nháp → gửi', 'Manager'],
        ['2', 'Manager', 'Mở Báo giá, bấm "Duyệt"', 'Nếu < ngưỡng → Đã duyệt; nếu ≥ ngưỡng → Chờ CEO', 'CEO (nếu lớn)'],
        ['3', 'CEO', 'Bấm "Duyệt cấp 2"', 'Báo giá: Đã duyệt; CEO nhận thông báo 🔔 ở bước 2', 'Sale'],
        ['4', 'Sale', 'Báo giá cho khách hàng (ngoài hệ thống)', 'Khách xem xét', 'Khách hàng'],
        ['5', 'Khách', 'Đồng ý mua', '—', 'Sale'],
        ['6', 'Sale', 'Bấm "Tạo HĐ" và/hoặc "Tạo đơn" trên báo giá đã duyệt', 'Sinh Hợp đồng + Đơn bán (Nháp)', 'Manager'],
        ['7', 'Manager', 'Mở đơn bán, bấm "Ký"', 'Đơn: Hiệu lực (active)', 'Kho'],
        ['8', 'Manager/Sale', 'Bấm "Giao" trên đơn', 'Đơn: Đang giao; TỰ SINH phiếu xuất WMS', 'Nhân viên kho'],
        ['9', 'NV kho', 'Mở phiếu xuất → "Quét soạn" (mã + ô + SL)', 'Trừ tồn từng ô; đủ → Đã soạn', 'NV kho'],
        ['10', 'NV kho', 'Bấm "Giao hàng"', 'Phiếu xuất: Đã giao; tồn đã trừ; serial→đã bán', 'Manager'],
        ['11', 'Manager', 'Ghi nhận thanh toán', 'Đã thu tăng → Công nợ giảm', 'Kế toán/CEO'],
        ['12', 'CEO', 'Xem báo cáo (doanh thu/công nợ tự cập nhật)', 'Dashboard phản ánh đơn vừa bán', '—'],
    ])
    kv('Lưu ý duyệt 2 cấp', 'Báo giá nhỏ (< ngưỡng, mặc định 100 triệu) chỉ cần Manager duyệt 1 lần. '
       'Báo giá lớn (≥ ngưỡng) thì sau Manager (cấp 1) phải có CEO duyệt (cấp 2) mới thành "Đã duyệt".')
    kv('Chống tự duyệt', 'Người tạo không tự duyệt báo giá của mình; người duyệt cấp 1 ≠ cấp 2 (trừ admin).')

    h('D.2. Luồng NHẬP HÀNG vào kho', 2)
    flow("NCC giao hàng ─► KHO tạo phiếu nhập ─► Quét nhận từng mã ─► Xác nhận ─► CỘNG TỒN")
    scenario('Kịch bản', [
        ['1', 'NV kho', 'Tạo đơn nhập (chọn kho + dòng hàng + ô đích + lô/hạn dùng nếu có)', 'Phiếu nhập: Nháp', 'NV kho'],
        ['2', 'NV kho', 'Bấm "Quét" → quét mã + nhập SL từng mặt hàng', 'Cộng dồn "đã nhận / cần nhận"', 'NV kho'],
        ['3', 'NV kho', 'Khi đủ, bấm "Xác nhận nhận"', 'CỘNG TỒN vào ô; tạo Lô (FEFO); ghi Lịch sử kho', '—'],
    ])

    h('D.3. Luồng KIỂM KÊ kho', 2)
    flow("NV kho tạo phiên + quét đếm ─► QUẢN LÝ KHO xem chênh lệch ─► Áp dụng (điều chỉnh tồn)")
    scenario('Kịch bản', [
        ['1', 'NV kho', 'WMS → Kiểm kê → "Phiên mới"', 'Phiên: Đang đếm', 'NV kho'],
        ['2', 'NV kho', 'Quét đếm từng ô (mã + ô + số đếm)', 'Lưu: tồn hệ thống vs số đếm', 'Quản lý kho'],
        ['3', 'Quản lý kho', 'Xem bảng chênh lệch (xanh dư / đỏ thiếu)', '—', 'Quản lý kho'],
        ['4', 'Quản lý kho', 'Bấm "Áp dụng"', 'Tồn = số đếm; ghi Lịch sử kho; phiên: Đã áp dụng', '—'],
    ])
    kv('Phân quyền', 'Nhân viên kho ĐẾM được nhưng KHÔNG duyệt; chỉ Quản lý kho trở lên bấm "Áp dụng".')

    h('D.4. Luồng LEAD → KHÁCH HÀNG', 2)
    scenario('Kịch bản', [
        ['1', 'Sale', 'Tạo Lead (tên/công ty/SĐT/nguồn)', 'Lead: new', 'Sale'],
        ['2', 'Sale', 'Liên hệ, cập nhật trạng thái', 'contacted → qualified', 'Sale'],
        ['3', 'Sale', 'Bấm "Chuyển KH"', 'Tạo Khách hàng thật; Lead: converted', 'Sale (tạo cơ hội)'],
    ])

    h('D.5. Luồng DỊCH VỤ / BẢO HÀNH', 2)
    scenario('Kịch bản', [
        ['1', 'Khách/Sale', 'Báo sản phẩm lỗi', '—', 'Dịch vụ'],
        ['2', 'Dịch vụ', 'Tạo Ticket, gắn số serial', 'Ticket: open', 'Dịch vụ'],
        ['3', 'Dịch vụ', 'Tra "Lịch sử serial" → biết đã bán cho ai, còn bảo hành?', '—', 'Dịch vụ'],
        ['4', 'Dịch vụ', 'Xử lý → cập nhật trạng thái', 'in_progress → resolved → closed', '—'],
    ])

    # ════════ E. TRỢ LÝ + THÔNG BÁO ════════
    h('E. Công cụ hỗ trợ', 1)
    kv('Trợ lý nội bộ (chat đáy trang)', 'Gõ lệnh để làm nhanh theo quyền: '
       '"làm báo giá cho ABC: 5 x 001002", "tồn 001002", "báo cáo điều hành". '
       'Hành động tạo BẢN NHÁP để người dùng xác nhận.')
    kv('Thông báo (chuông 🔔)', 'Báo khi báo giá chờ CEO duyệt / được duyệt / bị từ chối. '
       'Bấm để tới màn liên quan.')

    # ════════ F. PHÂN QUYỀN ════════
    h('F. Phân quyền tóm tắt', 1)
    table(['Chức năng', 'Sale', 'NV kho', 'QL kho', 'Manager', 'CEO/Admin'],
          [['Tạo/sửa báo giá, hợp đồng', '✓', '–', '–', '✓', '✓'],
           ['Duyệt báo giá cấp 1', '–', '–', '–', '✓', '✓'],
           ['Duyệt báo giá cấp 2', '–', '–', '–', '–', '✓'],
           ['Nhập/xuất/quét kho', '–', '✓', '✓', '✓', '✓'],
           ['Điều chỉnh tồn / duyệt kiểm kê', '–', '✗', '✓', '✓', '✓'],
           ['KPI vận hành kho', '–', '–', '✓', '✓', '✓'],
           ['Dashboard CEO / tài chính', '–', '–', '–', '✓', '✓']])

    # ════════ G. HƯỚNG DẪN CHI TIẾT TỪNG TRANG ════════
    d.add_page_break()
    h('G. Hướng dẫn chi tiết từng trang', 1)
    para('Tra cứu nhanh: mỗi trang gồm đường dẫn menu, các nút/trường và các bước thao tác.')

    h('G.0. Thanh chung (mọi trang)', 2)
    buttons([
        ('🔔 Chuông', 'Số thông báo chưa đọc; bấm xem danh sách; "Đọc hết" để xoá badge.'),
        ('Tên user + vai trò', 'Hiển thị người đang đăng nhập.'),
        ('Đăng xuất', 'Thoát tài khoản.'),
        ('Bộ chuyển CRM/WMS/CEO', 'Đổi phân hệ (chỉ hiện phân hệ bạn có quyền).'),
        ('Trợ lý nội bộ (đáy)', 'Gõ lệnh + "Hỏi"; chip gợi ý; Thu gọn; 🗑 xoá hội thoại.'),
    ])

    # ── CRM pages ──
    h('G.1. Khách hàng  (CRM → Khách hàng)', 2)
    buttons([
        ('Ô tìm kiếm', 'Lọc theo tên/mã KH.'),
        ('Import (Quản lý+)', 'Nhập KH cũ từ Excel/CSV.'),
        ('Thêm KH', 'Mở form tạo khách hàng.'),
        ('Bấm 1 dòng', 'Mở hồ sơ Khách hàng 360.'),
        ('Trước / Sau', 'Chuyển trang.'),
    ])
    para('Thêm khách hàng:')
    steps(['Bấm "Thêm KH".', 'Nhập Mã KH (bắt đầu "KH"), Tên, Phân khúc, Vùng, MST.',
           'Bấm "Thêm người liên hệ" để thêm liên hệ (đánh dấu "chính").', 'Bấm "Lưu".'])

    h('G.2. Khách hàng 360', 2)
    buttons([
        ('← Quay lại', 'Về danh sách.'),
        ('Sửa', 'Sửa thông tin KH.'),
        ('Thẻ KPI', 'Đơn mở, công nợ, ticket, hoạt động gần nhất.'),
        ('Lịch sử làm việc', 'Dòng thời gian tương tác; 🎧 nghe ghi âm, 📄 tải recap.'),
    ])

    h('G.3. Leads  (CRM → Leads)', 2)
    buttons([
        ('Tạo Lead', 'Tạo khách tiềm năng.'),
        ('Chuyển KH', 'Biến lead đủ điều kiện thành Khách hàng.'),
        ('Import (Quản lý+)', 'Nhập lead cũ.'),
        ('Bấm 1 dòng', 'Sửa lead.'),
    ])

    h('G.4. Cơ hội  (CRM → Opportunity / Pipeline)', 2)
    buttons([
        ('Tạo Opportunity', 'Tạo cơ hội (KH, giá trị, xác suất, giai đoạn).'),
        ('Bấm 1 dòng', 'Mở chi tiết + dòng thời gian; đổi giai đoạn.'),
        ('Pipeline (menu)', 'Xem kanban theo giai đoạn.'),
    ])

    h('G.5. Báo giá  (CRM → Báo giá)', 2)
    buttons([
        ('Tạo BG', 'Mở form: chọn KH, thêm dòng (mã/tên/SL/đơn giá); tổng tự tính.'),
        ('Bấm dòng (Nháp)', 'Sửa báo giá (chỉ khi Nháp).'),
        ('Duyệt / Duyệt (cấp 1)', 'Manager+ duyệt; <ngưỡng→Đã duyệt, ≥ngưỡng→Chờ CEO.'),
        ('Duyệt cấp 2 (CEO)', 'CEO/Admin duyệt báo giá Chờ CEO.'),
        ('Từ chối', 'Manager+ từ chối kèm lý do.'),
        ('Tạo đơn', 'Báo giá đã duyệt → tạo Đơn bán.'),
        ('Tạo HĐ', 'Báo giá đã duyệt → tạo Hợp đồng.'),
    ])
    para('Tạo & trình duyệt báo giá:')
    steps(['Bấm "Tạo BG", chọn khách hàng.', 'Nhập từng dòng: Mã part, Tên, SL, Đơn giá.',
           'Bấm "Tạo" (lưu Nháp).', 'Manager mở báo giá bấm "Duyệt".',
           'Nếu lớn: CEO bấm "Duyệt cấp 2".', 'Khi "Đã duyệt": bấm "Tạo đơn"/"Tạo HĐ".'])

    h('G.6. Hợp đồng / Công nợ', 2)
    buttons([
        ('Tạo HĐ', 'Tạo hợp đồng (hoặc sinh từ báo giá).'),
        ('Import (HĐ/Đơn cũ)', 'Nhập dữ liệu cũ; cột customer_code = mã KH.'),
        ('Phân tích tuổi nợ', '(Công nợ) Thanh tỷ lệ theo nhóm tuổi nợ.'),
    ])

    h('G.7. Viếng thăm / Hoạt động', 2)
    buttons([
        ('Khách hàng* / Ngày*', 'Bắt buộc.'),
        ('Mục đích / Nội dung / Tóm tắt / Bước tiếp', 'Nội dung buổi gặp/gọi.'),
        ('File ghi âm — Chọn file', 'Tải audio lên.'),
        ('File recap — Chọn file', 'Tải file recap (Word/PDF).'),
        ('Recap (văn bản)', 'Gõ tóm tắt.'),
        ('Lưu', 'Lưu; tự hiện ở Lịch sử làm việc của KH.'),
    ])

    h('G.8. Import dữ liệu (hộp thoại)', 2)
    buttons([
        ('Tải file mẫu (Excel)', 'Tải file mẫu đúng cột.'),
        ('Chọn file', 'Chọn .xlsx/.csv đã điền.'),
        ('Xem trước', 'Kiểm lỗi (sẽ tạo/bỏ qua trùng/lỗi từng dòng), CHƯA ghi.'),
        ('Import', 'Ghi vào hệ thống; bỏ qua bản ghi trùng mã.'),
    ])

    # ── WMS pages ──
    h('G.9. Nhập kho  (WMS → Nhập kho)', 2)
    buttons([
        ('Tạo đơn nhập', 'Chọn kho + dòng hàng + ô đích (+ lô/hạn dùng).'),
        ('Quét', 'Mở cửa sổ quét nhận theo phiếu.'),
        ('Xác nhận', 'Cộng tồn theo số đã nhận.'),
    ])
    steps(['Tạo đơn nhập.', 'Bấm "Quét" → quét/nhập mã + SL từng mặt hàng.',
           'Khi đủ ("Đủ"), bấm "Xác nhận nhận" → cộng tồn + tạo Lô.'])

    h('G.10. Xuất kho  (WMS → Xuất kho)', 2)
    buttons([
        ('Tạo đơn xuất', 'Tạo phiếu (hoặc auto khi Giao đơn bán).'),
        ('Pick-list', 'Xem gợi ý lấy hàng ở ô nào (FIFO/FEFO).'),
        ('Quét', 'Cửa sổ quét soạn: mã + ô + SL → trừ tồn.'),
        ('Giao', 'Xác nhận giao → trừ tồn, ghi xuất kho.'),
    ])

    h('G.11. Quét mã  (WMS → Quét mã)', 2)
    buttons([
        ('Tab Tra cứu', 'Quét → xem phụ tùng/serial.'),
        ('Tab Nhập kho', 'Quét mã + ô + SL → cộng tồn.'),
        ('Tab Xuất kho', 'Quét mã + ô + SL → trừ tồn.'),
        ('Tab Kiểm kê (Quản lý kho)', 'Quét mã + ô + số đếm → đặt lại tồn.'),
        ('Bắt đầu quét / Dừng', 'Bật/tắt camera.'),
        ('Đang quét vào: Mã hàng / Ô', 'Chọn camera điền vào trường nào.'),
    ])

    h('G.12. Kiểm kê  (WMS → Kiểm kê)', 2)
    buttons([
        ('Phiên mới', 'Tạo phiên kiểm kê (NV kho làm được).'),
        ('Mã hàng / Mã ô / Số đếm + quét', 'Ghi số đếm thực tế từng ô.'),
        ('Bảng chênh lệch', 'Hệ thống vs Đếm (xanh dư / đỏ thiếu).'),
        ('Áp dụng (Quản lý kho)', 'Điều chỉnh tồn theo số đếm.'),
    ])

    h('G.13. Lô hàng / Serial', 2)
    buttons([
        ('Lô hàng — Sắp hết hạn (≤30 ngày)', 'Lọc lô gần hết hạn (cảnh báo đỏ/vàng).'),
        ('Serial (menu)', 'Danh sách serial; tra lịch sử (đã bán cho ai, bảo hành).'),
    ])

    h('G.14. KPI vận hành  (WMS → KPI vận hành, Quản lý kho+)', 2)
    buttons([
        ('Chọn kỳ 7/30/90 ngày', 'Khoảng thời gian thống kê.'),
        ('Thẻ KPI', 'Nhập/Xuất (SL+lần), Độ chính xác kiểm kê, Sắp hết hàng.'),
        ('Bảng Tồn theo zone', 'SKU + tồn từng zone.'),
        ('Bảng Hiệu suất nhân sự', 'Số thao tác theo người.'),
    ])

    h('G.15. CEO  (menu CEO, Manager/CEO/Admin)', 2)
    buttons([
        ('Bảng điều hành', 'KPI tổng hợp công ty.'),
        ('AI Summary — Làm mới', 'Tổng hợp lại số liệu.'),
        ('AI Summary — Tải Excel', 'Xuất báo cáo điều hành .xlsx.'),
        ('Doanh thu / Công nợ / Forecast / Tồn', 'Báo cáo chuyên đề.'),
    ])

    d.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    main(sys.argv[1] if len(sys.argv) > 1 else 'HUONG_DAN_SU_DUNG.docx')
